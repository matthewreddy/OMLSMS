"""This file defines the logic for the dialog box shown if Bill is pressed in RenewalDlg."""

import datetime, os

from PyQt5.QtCore import *
from PyQt5.QtGui import *
from PyQt5.QtSql import *
from PyQt5.QtWidgets import *

import ui.ui_bill as ui
from constants import *
from formviewdlg import MainDlg

import djprint

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class BillDlg(MainDlg, ui.Ui_billDlg):
    def __init__(self, id, parent = None):
        super(BillDlg, self).__init__(parent)
        self.setupUi(self)
        self.setWindowFlags(Qt.Dialog | Qt.FramelessWindowHint | Qt.WindowTitleHint)
        self.setWindowTitle("OMLSMS Billing")
        self.setStyleSheet("QDialog { border: 1px solid black; }")
        self.sterilizer_id = id
    


    def html_to_invoice_docx(self, html: str, filename: str = "bill.docx"):
        """ This function is designed to take in html from djprint.getBillforSterilizer 
        and return a word document in similar format. """


        soup = BeautifulSoup(html, 'html.parser')
        doc = Document()

        # Title
        header_text = soup.find('td', class_='invoice')
        if header_text:
            p = doc.add_paragraph(header_text.get_text(strip=True))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(18)

        # Address block
        addr_lines = soup.find_all('td', class_='address_dentist')
        for td in addr_lines:
            doc.add_paragraph(td.get_text(strip=True))

        # Charges table
        table_html = soup.find_all('table', class_='charges')
        for table_tag in table_html:
            doc_table = doc.add_table(rows=0, cols=0)
            first_row = table_tag.find('tr')
            headers = first_row.find_all(['th', 'td'])

            doc_table.add_row()
            for _ in headers:
                doc_table.add_column(width=Pt(60))

            hdr_cells = doc_table.rows[0].cells
            for i, th in enumerate(headers):
                hdr_cells[i].text = th.get_text(strip=True)

            for tr in table_tag.find_all('tr')[1:]:
                cols = tr.find_all(['td', 'th'])
                row_cells = doc_table.add_row().cells
                for i, td in enumerate(cols):
                    row_cells[i].text = td.get_text(strip=True)

            doc.add_paragraph()  # spacing between tables

        # Late payment notice
        notices = soup.find_all('td', class_='late_notice')
        for n in notices:
            p = doc.add_paragraph(n.get_text(strip=True))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.runs[0].bold = True

        # Divider (optional)
        dividers = soup.find_all('td', class_='divider')
        for d in dividers:
            p = doc.add_paragraph(d.get_text(strip=True))
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.runs[0].italic = True

        # Footer message
        footer = soup.find('div', id='footerContent')
        if footer:
            lines = footer.find_all('td', class_='invoice')
            for line in lines:
                p = doc.add_paragraph(line.get_text(strip=True))
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Save doc
        doc.save(filename)


    # Functions that defined behavior for buttons when clicked
    @pyqtSlot()
    def on_defaultPrintButton_clicked(self) -> None:
        self.printHTML(djprint.getBillForSterilizer(self.sterilizer_id))
        self.close()
    
    @pyqtSlot()
    def on_wordDocButton_clicked(self) -> None:
        today = datetime.datetime.today().date().isoformat()
        path = str(self.sterilizer_id) + " Bill " + today + ".docx"
        self.html_to_invoice_docx(djprint.getBillForSterilizer(self.sterilizer_id), path)
        os.startfile(path)
        
        self.close()

    @pyqtSlot()
    def on_cancelPushButton_clicked(self) -> None:
        self.close()

